"""Utilities for parsing Word documents into theory outline drafts."""

from __future__ import annotations

import re
from dataclasses import dataclass
from html import escape
from io import BytesIO
from typing import List, Optional

from docx import Document  # type: ignore[import-untyped]


@dataclass
class LessonDraft:
    """Lightweight representation of a lesson parsed from a Word document."""

    title: str
    content_html: str
    summary: str
    order_index: int


@dataclass
class TopicDraft:
    """Lightweight representation of a theory topic parsed from Word."""

    title: str
    summary: str
    intro_html: str
    lessons: List[LessonDraft]
    order_index: int


@dataclass
class ChapterDraft:
    """Lightweight representation of a theory chapter parsed from Word."""

    title: str
    summary: str
    intro_html: str
    topics: List[TopicDraft]
    order_index: int


def _normalize_style_name(style_name: Optional[str]) -> str:
    if not style_name:
        return ""
    return style_name.strip().lower()


def _detect_heading_level(style_name: Optional[str]) -> int:
    """Return heading level (1-3) if the style represents a heading, else 0."""

    normalized = _normalize_style_name(style_name)
    if not normalized:
        return 0

    match = re.search(r"(heading|标题)\s*(\d)", normalized)
    if match:
        return int(match.group(2))

    fallback_map = {
        "heading1": 1,
        "heading 1": 1,
        "标题1": 1,
        "标题 1": 1,
        "heading2": 2,
        "heading 2": 2,
        "标题2": 2,
        "标题 2": 2,
        "heading3": 3,
        "heading 3": 3,
        "标题3": 3,
        "标题 3": 3,
    }
    return fallback_map.get(normalized, 0)


_NUMBER_PREFIX_RE = re.compile(
    r"^\s*(?:\d+(?:[.．]\d+)*|[一二三四五六七八九十百千零〇]+)\s*(?:[.)、．）])?\s*"
)


def _clean_heading_text(text: str) -> str:
    if not text:
        return ""
    stripped = text.strip()
    stripped = _NUMBER_PREFIX_RE.sub("", stripped, count=1)
    stripped = re.sub(r"^第[一二三四五六七八九十百零〇0-9]+[章节篇部分]\s*", "", stripped)
    return stripped.strip()


def _paragraph_runs_to_html(paragraph) -> str:
    fragments: List[str] = []
    for run in paragraph.runs:
        text = escape(run.text or "")
        if not text:
            continue
        if getattr(run, "hyperlink", None) and getattr(run.hyperlink, "target", None):
            href = escape(run.hyperlink.target)
            text = f"<a href=\"{href}\" target=\"_blank\" rel=\"noopener noreferrer\">{text}</a>"
        if run.bold:
            text = f"<strong>{text}</strong>"
        if run.italic:
            text = f"<em>{text}</em>"
        if run.underline:
            text = f"<u>{text}</u>"
        fragments.append(text)
    combined = "".join(fragments)
    if not combined:
        return ""
    return f"<p>{combined}</p>"


def _strip_html(html: str) -> str:
    if not html:
        return ""
    text = re.sub(r"<[^>]+>", " ", html)
    return re.sub(r"\s+", " ", text).strip()


def _summarize_html(html: str, limit: int = 120) -> str:
    text = _strip_html(html)
    if len(text) <= limit:
        return text
    return text[: limit - 1].rstrip() + "…"


def _finalize_lesson(
    current_lesson: Optional[LessonDraft],
    lesson_body: List[str],
    lessons: List[LessonDraft],
) -> None:
    if current_lesson is None:
        return
    html = "".join(lesson_body).strip()
    if not html:
        html = "<p><br></p>"
    current_lesson.content_html = html
    current_lesson.summary = _summarize_html(html)
    lessons.append(current_lesson)


def _finalize_topic(
    current_topic: Optional[TopicDraft],
    intro_fragments: List[str],
    current_chapter: Optional[ChapterDraft],
) -> None:
    if current_topic is None or current_chapter is None:
        return
    intro_html = "".join(intro_fragments).strip()
    current_topic.intro_html = intro_html
    current_topic.summary = _summarize_html(intro_html)
    current_chapter.topics.append(current_topic)


def _finalize_chapter(
    current_chapter: Optional[ChapterDraft],
    intro_fragments: List[str],
    chapters: List[ChapterDraft],
) -> None:
    if current_chapter is None:
        return
    intro_html = "".join(intro_fragments).strip()
    current_chapter.intro_html = intro_html
    current_chapter.summary = _summarize_html(intro_html)
    chapters.append(current_chapter)


def _ensure_chapter(chapters: List[ChapterDraft]) -> ChapterDraft:
    title = f"自动生成章节 {len(chapters) + 1}"
    return ChapterDraft(title=title, summary="", intro_html="", topics=[], order_index=len(chapters))


def parse_docx_outline(file_storage) -> dict:
    """Parse a .docx file and return a structured outline for import previews."""

    if not file_storage:
        raise ValueError("No file provided")

    # Ensure the stream pointer is at the start before python-docx reads it.
    raw = getattr(file_storage, "stream", None)
    if raw is None:
        raw = file_storage
    else:
        raw.seek(0)

    # python-docx expects a path or a binary file-like object.
    if not hasattr(raw, "read"):
        buffer = BytesIO(file_storage.read())  # type: ignore[arg-type]
        buffer.seek(0)
        document = Document(buffer)
    else:
        document = Document(raw)

    chapters: List[ChapterDraft] = []
    warnings: List[str] = []

    current_chapter: Optional[ChapterDraft] = None
    current_topic: Optional[TopicDraft] = None
    current_lesson: Optional[LessonDraft] = None
    chapter_intro_fragments: List[str] = []
    topic_intro_fragments: List[str] = []
    lesson_body: List[str] = []

    for paragraph in document.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        level = _detect_heading_level(style_name)
        raw_text = paragraph.text.strip()
        heading_text = _clean_heading_text(raw_text)

        if level == 1:
            if current_topic is not None:
                _finalize_lesson(current_lesson, lesson_body, current_topic.lessons)
            current_lesson = None
            lesson_body = []
            _finalize_topic(current_topic, topic_intro_fragments, current_chapter)
            current_topic = None
            topic_intro_fragments = []
            _finalize_chapter(current_chapter, chapter_intro_fragments, chapters)
            chapter_title = heading_text or raw_text or f"未命名章节 {len(chapters) + 1}"
            current_chapter = ChapterDraft(
                title=chapter_title,
                summary="",
                intro_html="",
                topics=[],
                order_index=len(chapters),
            )
            chapter_intro_fragments = []
            continue

        if level == 2:
            if current_chapter is None:
                warnings.append("检测到二级标题前没有一级标题，已为其创建默认章节。")
                current_chapter = _ensure_chapter(chapters)
                chapter_intro_fragments = []
            if current_topic is not None:
                _finalize_lesson(current_lesson, lesson_body, current_topic.lessons)
                current_lesson = None
                lesson_body = []
                _finalize_topic(current_topic, topic_intro_fragments, current_chapter)
                topic_intro_fragments = []
            topic_title = heading_text or raw_text or f"未命名目录 {len(current_chapter.topics) + 1}"
            current_topic = TopicDraft(
                title=topic_title,
                summary="",
                intro_html="",
                lessons=[],
                order_index=len(current_chapter.topics),
            )
            topic_intro_fragments = []
            continue

        if level == 3:
            if current_topic is None:
                warnings.append("检测到三级标题前没有二级标题，将忽略该标题。")
                continue
            _finalize_lesson(current_lesson, lesson_body, current_topic.lessons)
            lesson_title = heading_text or raw_text or f"未命名知识点 {len(current_topic.lessons) + 1}"
            current_lesson = LessonDraft(
                title=lesson_title,
                content_html="",
                summary="",
                order_index=len(current_topic.lessons),
            )
            lesson_body = []
            continue

        html = _paragraph_runs_to_html(paragraph)
        if not html:
            continue
        if current_lesson is not None:
            lesson_body.append(html)
        elif current_topic is not None:
            topic_intro_fragments.append(html)
        elif current_chapter is not None:
            chapter_intro_fragments.append(html)
        else:
            warnings.append("文档开始部分缺少一级标题，已创建默认章节以承载正文。")
            current_chapter = _ensure_chapter(chapters)
            chapter_intro_fragments = [html]

    if current_topic is not None:
        _finalize_lesson(current_lesson, lesson_body, current_topic.lessons)
        _finalize_topic(current_topic, topic_intro_fragments, current_chapter)

    if current_chapter is not None:
        _finalize_chapter(current_chapter, chapter_intro_fragments, chapters)

    chapter_dicts = []
    topic_count = 0
    lesson_count = 0
    for chapter in chapters:
        topic_dicts = []
        for topic in chapter.topics:
            lesson_dicts = []
            for lesson in topic.lessons:
                lesson_dicts.append(
                    {
                        "title": lesson.title,
                        "contentHtml": lesson.content_html,
                        "summary": lesson.summary,
                        "orderIndex": lesson.order_index,
                    }
                )
            lesson_count += len(lesson_dicts)
            topic_dicts.append(
                {
                    "title": topic.title,
                    "summary": topic.summary,
                    "introHtml": topic.intro_html,
                    "lessons": lesson_dicts,
                    "orderIndex": topic.order_index,
                }
            )
        topic_count += len(topic_dicts)
        chapter_dicts.append(
            {
                "title": chapter.title,
                "summary": chapter.summary,
                "introHtml": chapter.intro_html,
                "topics": topic_dicts,
                "orderIndex": chapter.order_index,
            }
        )

    return {
        "fileName": getattr(file_storage, "filename", ""),
        "chapters": chapter_dicts,
        "warnings": warnings,
        "stats": {
            "chapterCount": len(chapter_dicts),
            "topicCount": topic_count,
            "lessonCount": lesson_count,
        },
    }
